import os ,hashlib
import csv
from utils.logger_handler import logger
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader,PyPDFLoader,UnstructuredMarkdownLoader,UnstructuredWordDocumentLoader
#获取文件的md5的十六进制字符串
def get_file_md5_hex(filepath:str):
    if not os.path.exists(filepath):
        logger.error(f"[md5计算]文件{filepath}不存在")
        return 
    
    if not os.path.isfile(filepath):
        logger.error(f"[md5计算]文件{filepath}不是文件")
        return 
    
    md5_obj = hashlib.md5()

    chunk_size = 4096 #4KB分片，避免文件内存过大

    try:
        with open(filepath,"rb")as f: #计算文件的md5 必须二进制读取rb
            while chunk :=f.read(chunk_size):
                md5_obj.update(chunk)
            """
            chunk = f.read(chunk_size)
            while chunk:
                md5_obj.update(chunk)
                chunk = f.read(chunk_size)
            """
            md5_hex = md5_obj.hexdigest()
            return md5_hex
    except Exception as e:
        logger.error(f"[md5计算]文件{filepath}计算md5时发生错误: {str(e)}")
        return




#返回文件夹内的文件列表，过滤掉非文本文件
def listdir_with_allowed_type(path:str,allowed_types:tuple[str]):
    files = []
    
    if not os.path.isdir(path):
        logger.error(f"[listdir_with_allowed_type]{path}不是文件夹")
        return allowed_types
    for f in os.listdir(path):
        if f.endswith(allowed_types):
            files.append(os.path.join(path,f))
    return tuple(files)

def pdf_loader(filepath:str,passdw = None)->list[Document]:
    return PyPDFLoader(filepath,passdw).load()

def txt_loader(filepath:str,encoding:str="utf-8")->list[Document]:
    return TextLoader(filepath,encoding).load()

def word_loader(filepath:str)->list[Document]:
    """加载Word文档"""
    try:
        return UnstructuredWordDocumentLoader(filepath).load()
    except Exception as e:
        logger.error(f"[word_loader]加载Word文档失败: {str(e)}")
        # 尝试手动解析
        docs = []
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            for table in doc.tables:
                table_text = "\n表格:\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text += row_text + "\n"
                content.append(table_text)
            full_content = "\n".join(content)
            if full_content.strip():
                docs.append(Document(
                    page_content=full_content,
                    metadata={"source": filepath}
                ))
        except Exception as e2:
            logger.error(f"[word_loader]手动解析Word文档失败: {str(e2)}")
        return docs

def markdown_loader(filepath:str,encoding:str="utf-8")->list[Document]:
    """加载Markdown文档"""
    try:
        return UnstructuredMarkdownLoader(filepath,encoding=encoding).load()
    except Exception as e:
        logger.error(f"[markdown_loader]加载Markdown文档失败: {str(e)}")
        # 尝试手动解析
        docs = []
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            if content.strip():
                docs.append(Document(
                    page_content=content,
                    metadata={"source": filepath}
                ))
        except Exception as e2:
            logger.error(f"[markdown_loader]手动解析Markdown文档失败: {str(e2)}")
        return docs

def csv_loader(filepath:str,encoding:str="utf-8")->list[Document]:
    """加载CSV文件"""
    docs = []
    try:
        with open(filepath, 'r', encoding=encoding) as f:
            reader = csv.reader(f)
            rows = list(reader)
            if rows:
                # 提取表头
                header = rows[0]
                content = "CSV文件内容:\n\n"
                content += "表头: " + ", ".join(header) + "\n\n"
                content += "数据:\n"
                for i, row in enumerate(rows[1:], 1):
                    content += f"行{i}: " + ", ".join(row) + "\n"
                docs.append(Document(
                    page_content=content,
                    metadata={"source": filepath}
                ))
    except Exception as e:
        logger.error(f"[csv_loader]加载CSV文件失败: {str(e)}")
    return docs
    